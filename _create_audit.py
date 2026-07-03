from docx import Document
from docx.shared import Pt, Cm

doc = Document()
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(11)

doc.add_heading('Abkürzungen — Anmerkungen & Spell-out-Audit', level=1)

# Section 1: omitted
doc.add_heading('Bewusst weggelassen (Entscheidung, anpassbar)', level=2)
omitted = [
    ("BYD", "funktioniert als Firmenname; „Build Your Dreams\" ist ein Marketing-Backronym, das akademisch nicht behauptet werden sollte."),
    ("EBIT, CCB", "je nur 1x im Text; CCB (China Construction Bank) ist im Text gar nicht erklärt."),
    ("HKEX", "kommt nur im Literaturverzeichnis vor, nicht im Fließtext."),
]
for abbr, reason in omitted:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f"{abbr} — ")
    r.bold = True
    p.add_run(reason)

# Section 2: spell-out audit
doc.add_heading('Spell-out-Audit — bei Erstnennung NICHT ausgeschrieben', level=2)
p = doc.add_paragraph()
p.add_run("Nur ").italic = False
run = p.add_run("CMA")
run.bold = True
p.add_run(" und ")
run = p.add_run("ZGH")
run.bold = True
p.add_run(" werden korrekt eingeführt. Folgende nicht:")

audit = [
    ("FDI", "4.3", "strategic asset-seeking foreign direct investment (FDI)"),
    ("MEB", "4.1", "Volkswagen's Modular Electric Drive Matrix (MEB) plant"),
    ("NEV", "2.2", "global new energy vehicle (NEV) market"),
    ("ROE", "3.1", "weighted average return on equity (ROE)"),
    ("SASAC", "2.1", "Shanghai's State-owned Assets Supervision and Administration Commission (SASAC)"),
    ("SOE", "2.4", "its state-owned enterprise (SOE) ownership structure"),
    ("ICE", "5.3", "no internal combustion engine (ICE) segment"),
    ("OEM", "2.2", "Western original equipment manufacturers (OEMs)"),
    ("IPO", "2.1", "three-year post-IPO… → … post-initial public offering (IPO)"),
    ("EBIT", "5.3", "6.8% core EBIT (earnings before interest and taxes) margin"),
    ("CCB", "3.3", "alongside the China Construction Bank (CCB)"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(["Abk.", "Stelle", "sollte bei Erstnennung heißen"]):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

for abbr, loc, fix in audit:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = loc
    row[2].text = fix
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

for row in table.rows:
    row.cells[0].width = Cm(2)
    row.cells[1].width = Cm(2)
    row.cells[2].width = Cm(13)

p = doc.add_paragraph()
p.add_run("Konvention: ").bold = True
p.add_run("Vollform beim ersten Mal, danach nur Abkürzung. GM ist okay — „General Motors\" steht in 3.1 vor dem ersten „SAIC-GM\".")

try:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations-audit.docx")
    print("Saved to abbreviations-audit.docx")
except PermissionError:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations-audit-v2.docx")
    print("Locked - saved to abbreviations-audit-v2.docx")
