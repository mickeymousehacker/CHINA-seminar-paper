# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING

def art(pre, journal, vol, post):
    return [(pre, False), (journal, True), (", ", False), (vol, True), (post, False)]

def rep(pre, title, post):
    return [(pre, False), (title, True), (post, False)]

refs = [
    art("Allen, F., Qian, J., Shan, C., & Zhu, J. L. (2024). Dissecting the long-term performance of the Chinese stock market. ",
        "Journal of Finance", "79", "(2), 1121–1171. https://doi.org/10.1111/jofi.13312"),
    art("Bai, J., Barwick, P. J., Cao, S., & Li, S. (2025). Quid pro quo, knowledge spillovers, and industrial quality upgrading: Evidence from the Chinese auto industry. ",
        "American Economic Review", "115", "(11), 3825–3852. https://doi.org/10.1257/aer.20221501"),
    art("Balcet, G., Wang, H., & Richet, X. (2012). Geely: A trajectory of catching up and asset-seeking multinational growth. ",
        "International Journal of Automotive Technology and Management", "12", "(4), 360–375. https://doi.org/10.1504/IJATM.2012.051361"),
    art("Bortolotti, B., & Faccio, M. (2009). Government control of privatized firms. ",
        "Review of Financial Studies", "22", "(8), 2907–2939. https://doi.org/10.1093/rfs/hhn077"),
    rep("BYD Company Limited. (2025). ", "Annual report 2024",
        ". Hong Kong Stock Exchange (HKEX: 01211). https://www.bydglobal.com/en/InvestorAnnals.html"),
    rep("BYD Company Limited. (2026). ", "Annual report 2025",
        ". Hong Kong Stock Exchange (HKEX: 01211). https://www.bydglobal.com/en/InvestorAnnals.html"),
    rep("European Commission. (2024). ",
        "Commission Implementing Regulation (EU) 2024/2754 of 29 October 2024 imposing a definitive countervailing duty on imports of new battery electric vehicles originating in the People’s Republic of China",
        ". Official Journal of the European Union. https://eur-lex.europa.eu/eli/reg_impl/2024/2754/oj/eng"),
    art("Fan, J. P. H., Wong, T. J., & Zhang, T. (2007). Politically connected CEOs, corporate governance, and post-IPO performance of China’s newly partially privatized firms. ",
        "Journal of Financial Economics", "84", "(2), 330–357. https://doi.org/10.1016/j.jfineco.2006.03.008"),
    rep("Geely Automobile Holdings Limited. (2025a). ", "Annual report 2024",
        ". Hong Kong Stock Exchange (HKEX: 00175). https://www.hkexnews.hk/listedco/listconews/sehk/2025/0320/2025032001294.pdf"),
    rep("Geely Automobile Holdings Limited. (2025b). ", "Announcement of interim results for the six months ended 30 June 2025",
        ". Hong Kong Stock Exchange (HKEX: 00175). https://www1.hkexnews.hk/listedco/listconews/sehk/2025/0814/2025081400145.pdf"),
    rep("Geely Automobile Holdings Limited. (2026). ", "Annual report 2025",
        ". Hong Kong Stock Exchange (HKEX: 00175). https://www.geelyauto.com.hk/wp-content/uploads/2026/04/e00175_2025-annual-report.pdf"),
    art("Holmes, T. J., McGrattan, E. R., & Prescott, E. C. (2015). Quid pro quo: Technology capital transfers for market access in China. ",
        "Review of Economic Studies", "82", "(3), 1154–1193. https://doi.org/10.1093/restud/rdv008"),
    art("Howell, S. T. (2018). Joint ventures and technology adoption: A Chinese industrial policy that backfired. ",
        "Research Policy", "47", "(8), 1448–1462. https://doi.org/10.1016/j.respol.2018.04.021"),
    art("Inkpen, A. C., & Beamish, P. W. (1997). Knowledge, bargaining power, and the instability of international joint ventures. ",
        "Academy of Management Review", "22", "(1), 177–202. https://doi.org/10.5465/amr.1997.9707180263"),
    rep("International Energy Agency. (2026). ", "Global EV outlook 2026",
        ". IEA. https://www.iea.org/reports/global-ev-outlook-2026"),
    art("Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behavior, agency costs and ownership structure. ",
        "Journal of Financial Economics", "3", "(4), 305–360. https://doi.org/10.1016/0304-405X(76)90026-X"),
    art("Jiang, Z., & Xu, C. (2023). Policy incentives, government subsidies, and technological innovation in new energy vehicle enterprises: Evidence from China. ",
        "Energy Policy", "177", ", Article 113527. https://doi.org/10.1016/j.enpol.2023.113527"),
    art("Konda, P., Slepnikov, D., & Jin, J. (2022). From transaction to co-creation in Geely’s acquisition of Volvo Cars. ",
        "Asian Journal of Technology Innovation", "31", "(3), 556–580. https://doi.org/10.1080/19761597.2022.2133777"),
    art("Luo, Y., & Tung, R. L. (2007). International expansion of emerging market enterprises: A springboard perspective. ",
        "Journal of International Business Studies", "38", "(4), 481–498. https://doi.org/10.1057/palgrave.jibs.8400275"),
    art("Megginson, W. L., & Netter, J. M. (2001). From state to market: A survey of empirical studies on privatization. ",
        "Journal of Economic Literature", "39", "(2), 321–389. https://doi.org/10.1257/jel.39.2.321"),
    rep("SAIC Motor Corporation Limited. (2025). ", "Annual report 2024",
        ". Shanghai Stock Exchange. https://www.saicmotor.com/english/images/investor_relations/annual_report/2025/5/30/A82D5D04C86944B6A245ACFD8A90FB95.pdf"),
    rep("SAIC Motor Corporation Limited. (2026). ", "Annual report 2025",
        ". Shanghai Stock Exchange. https://www.saicmotor.com/english/images/investor_relations/annual_report/2026/4/21/0DFE68E4BCB74B01ACFAE92C344763FF.pdf"),
    art("Shleifer, A., & Vishny, R. W. (1994). Politicians and firms. ",
        "Quarterly Journal of Economics", "109", "(4), 995–1025. https://doi.org/10.2307/2118354"),
    art("Teece, D. J., Pisano, G., & Shuen, A. (1997). Dynamic capabilities and strategic management. ",
        "Strategic Management Journal", "18", "(7), 509–533. https://doi.org/10.1002/(SICI)1097-0266(199708)18:7<509::AID-SMJ882>3.0.CO;2-Z"),
    art("Wang, X., Zhao, W., & Ruet, J. (2022). Specialised vertical integration: The value-chain strategy of EV lithium-ion battery firms in China. ",
        "International Journal of Automotive Technology and Management", "22", "(2), 178–201. https://doi.org/10.1504/IJATM.2022.124377"),
    art("Whitfield, L., & Wuttke, T. (2026). China’s technological catch-up and leapfrogging in electric vehicles: A firm-level study of BYD and CATL. ",
        "Progress in Economic Geography", "4", ", Article 100054. https://doi.org/10.1016/j.peg.2025.100054"),
    art("Yuan, J.-Z., & Brasó Broggi, C. (2025). The metamorphosis of China’s automotive industry (1953–2001): Inward internationalisation, technological transfers and the making of a post-socialist market. ",
        "Business History", "67", "(1), 211–238. https://doi.org/10.1080/00076791.2023.2247366"),
    art("Zhang, Z.-Q., Tang, B.-J., Su, Z., & Zhang, Y. (2025). Founder control and breakthrough innovation: Evidence from high-tech firms in China. ",
        "China Economic Review", "94", ", Article 102572. https://doi.org/10.1016/j.chieco.2025.102572"),
    art("Zheng, Q., Noorderhaven, N., & Du, J. (2022). Making the unlikely marriage work: The integration process of Chinese-acquired foreign firms. ",
        "Journal of World Business", "57", "(2), Article 101305. https://doi.org/10.1016/j.jwb.2021.101305"),
]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

h = doc.add_paragraph()
hr = h.add_run("References")
hr.bold = True; hr.font.size = Pt(14); hr.font.name = "Times New Roman"
h.paragraph_format.space_after = Pt(12)

for entry in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-1.27)   # hanging indent
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    for text, ital in entry:
        r = p.add_run(text)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)
        r.italic = ital

out = r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\references-APA7-clean.docx"
try:
    doc.save(out)
    print("Saved:", out)
except PermissionError:
    out = out.replace(".docx", "-v2.docx")
    doc.save(out); print("Locked -> saved:", out)
print("Entries:", len(refs))
