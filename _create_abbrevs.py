from docx import Document
from docx.shared import Pt, Cm

doc = Document()

heading = doc.add_heading('List of Abbreviations', level=1)

abbreviations = [
    ("CMA", "Compact Modular Architecture"),
    ("EV", "Electric Vehicle"),
    ("FDI", "Foreign Direct Investment"),
    ("GM", "General Motors"),
    ("ICE", "Internal Combustion Engine"),
    ("JV", "Joint Venture"),
    ("M&A", "Mergers and Acquisitions"),
    ("MEB", "Modular Electric Drive Matrix"),
    ("NEV", "New Energy Vehicle"),
    ("OEM", "Original Equipment Manufacturer"),
    ("RMB", "Renminbi"),
    ("ROE", "Return on Equity"),
    ("SAIC", "Shanghai Automotive Industry Corporation"),
    ("SASAC", "State-owned Assets Supervision and Administration Commission"),
    ("SOE", "State-Owned Enterprise"),
    ("ZGH", "Zhejiang Geely Holding Group"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = "Abbreviation"
hdr[1].text = "Full Term"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

for abbr, full in abbreviations:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = full
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(11)

try:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations.docx")
    print("Saved 16-entry list to abbreviations.docx")
except PermissionError:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations-v2.docx")
    print("File locked - saved to abbreviations-v2.docx instead")
