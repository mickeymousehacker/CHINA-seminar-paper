from docx import Document
import re

doc = Document(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\China Strategies_term paper_draft_02 (2).docx")
paras = [p.text for p in doc.paragraphs]

# Use heading styles to find real body boundaries
body_start, ref_start = None, None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 1" and p.text.strip().startswith("1. Introduction"):
        body_start = i
    if p.style.name == "Heading 1" and p.text.strip() == "References":
        ref_start = i

print(f"Body para {body_start} .. References para {ref_start}\n")

targets = ["FDI", "MEB", "NEV", "ROE", "SASAC", "SOE", "ICE", "JV", "CMA",
           "M&A", "GM", "CCB", "EBIT", "EV", "IPO", "ZGH", "RMB", "USD", "SEK", "CNY"]

for ab in targets:
    pat = re.compile(r'\b' + re.escape(ab) + r'\b')
    found = False
    for i in range(body_start, ref_start):
        if pat.search(paras[i]):
            m = pat.search(paras[i])
            start = max(0, m.start()-120)
            snippet = paras[i][start:m.end()+12].replace("\n"," ")
            print(f"[{ab}] p{i}: ...{snippet}...")
            found = True
            break
    if not found:
        print(f"[{ab}] -- not in body --")
