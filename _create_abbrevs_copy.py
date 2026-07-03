from docx import Document
from docx.shared import Pt, Cm

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)

doc.add_heading('List of Abbreviations', level=1)

abbreviations = [
    ("CMA", "Compact Modular Architecture"),
    ("EV", "Electric Vehicle"),
    ("FDI", "Foreign Direct Investment"),
    ("GM", "General Motors"),
    ("ICE", "Internal Combustion Engine"),
    ("JV", "Joint Venture"),
    ("M&A", "Mergers and Acquisitions"),
    ("MEB", "Modular Electric Drive Matrix"),
    ("NEV", "New Energy Vehicle"),
    ("OEM", "Original Equipment Manufacturer"),
    ("RMB", "Renminbi"),
    ("ROE", "Return on Equity"),
    ("SAIC", "Shanghai Automotive Industry Corporation"),
    ("SASAC", "State-owned Assets Supervision and Administration Commission"),
    ("SOE", "State-Owned Enterprise"),
    ("ZGH", "Zhejiang Geely Holding Group"),
]

# Borderless table for clean copy-paste alignment
table = doc.add_table(rows=0, cols=2)
table.autofit = False
for abbr, full in abbreviations:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = full
    row[0].width = Cm(3.5)
    row[1].width = Cm(12)
    for cell in row:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

try:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations-copy.docx")
    print("Saved clean copyable version to abbreviations-copy.docx")
except PermissionError:
    doc.save(r"C:\Users\User\CHINA-seminar-paper\research\wiki\saic-motor\abbreviations-copy-v2.docx")
    print("Locked - saved to abbreviations-copy-v2.docx")
